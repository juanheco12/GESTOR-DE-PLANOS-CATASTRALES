from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime
from config import settings
import re


def _crear_template_defecto() -> Document:
    """Builds a basic Word template when no custom template is uploaded."""
    doc = Document()

    # ── Title ──────────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RESOLUCIÓN DE MUTACIÓN CATASTRAL")
    r.bold = True
    r.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("MUTACIÓN DE TERCERA CLASE — INCORPORACIÓN DE CONSTRUCCIÓN")
    r2.bold = True
    r2.font.size = Pt(11)

    doc.add_paragraph()

    # ── Header data ─────────────────────────────────────────────────────────
    for label, marker in [
        ("Expediente No.:", "{{EXPEDIENTE}}"),
        ("Código Catastral:", "{{NUMERO_PREDIO}}"),
        ("Propietario:", "{{PROPIETARIO}}"),
        ("Fecha de Solicitud:", "{{FECHA_SOLICITUD}}"),
        ("Inspector:", "{{INSPECTOR}}"),
    ]:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label} ")
        run_label.bold = True
        p.add_run(marker)

    doc.add_paragraph()

    # ── Separator ───────────────────────────────────────────────────────────
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.add_run("─" * 60)

    doc.add_paragraph()

    # ── Motivada body ───────────────────────────────────────────────────────
    body_title = doc.add_paragraph()
    body_title.add_run("MOTIVADA:").bold = True

    doc.add_paragraph("{{MOTIVADA_TEXTO}}")

    doc.add_paragraph()

    # ── Signature block ──────────────────────────────────────────────────────
    sig_label = doc.add_paragraph()
    sig_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_label.add_run("{{INSPECTOR}}").bold = True

    sig_title = doc.add_paragraph()
    sig_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_title.add_run("Inspector Catastral")

    fecha_exp = doc.add_paragraph()
    fecha_exp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha_exp.add_run(f"Fecha de expedición: {datetime.now().strftime('%d de %B de %Y')}")

    return doc


def _reemplazar_en_parrafo(paragraph, old: str, new: str):
    """Replace a marker in a paragraph preserving the run's style."""
    if old not in paragraph.text:
        return
    full_text = paragraph.text
    new_full = full_text.replace(old, new)
    # Clear all runs and put everything in the first run
    for i, run in enumerate(paragraph.runs):
        run.text = new_full if i == 0 else ""


def _inyectar_motivada(paragraph, motivada_texto: str):
    """Replace {{MOTIVADA_TEXTO}} marker with the full motivada body."""
    if "{{MOTIVADA_TEXTO}}" not in paragraph.text:
        return
    for run in paragraph.runs:
        if "{{MOTIVADA_TEXTO}}" in run.text:
            run.text = run.text.replace("{{MOTIVADA_TEXTO}}", motivada_texto)
            return
    # Fallback: clear and rewrite first run
    for i, run in enumerate(paragraph.runs):
        run.text = motivada_texto if i == 0 else ""


def generar_documento_word(motivada_texto: str, datos: dict, template_path: str = None) -> str:
    """Fill the Word template with generated motivada and form data. Returns filename."""

    doc = Document(template_path) if (template_path and Path(template_path).exists()) else _crear_template_defecto()

    propietario = datos.get("propietario", {})
    construccion = datos.get("construccion", {})

    prop = propietario if isinstance(propietario, dict) else {}
    con = construccion if isinstance(construccion, dict) else {}
    funcionario = f"{datos.get('funcionario_nombre', '')} — {datos.get('funcionario_cargo', '')}".strip(" —")

    reemplazos = {
        "{{EXPEDIENTE}}": datos.get("expediente", ""),
        "{{NUMERO_PREDIO}}": datos.get("numero_predio", ""),
        "{{PROPIETARIO}}": prop.get("nombre", ""),
        "{{FECHA_SOLICITUD}}": datos.get("fecha_solicitud", ""),
        "{{INSPECTOR}}": funcionario,
        "{{AREA}}": str(con.get("area_construida", "")),
        "{{DIRECCION}}": prop.get("direccion", ""),
        "{{MUNICIPIO}}": prop.get("municipio", ""),
        "{{FUNCIONARIO_NOMBRE}}": datos.get("funcionario_nombre", ""),
        "{{FUNCIONARIO_CARGO}}": datos.get("funcionario_cargo", ""),
    }

    for paragraph in doc.paragraphs:
        if "{{MOTIVADA_TEXTO}}" in paragraph.text:
            _inyectar_motivada(paragraph, motivada_texto)
        else:
            for old, new in reemplazos.items():
                _reemplazar_en_parrafo(paragraph, old, str(new))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old, new in reemplazos.items():
                        _reemplazar_en_parrafo(paragraph, old, str(new))

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    exp_clean = re.sub(r"[^\w\-]", "_", datos.get("expediente", "SIN_EXP"))
    filename = f"motivada_{exp_clean}_{timestamp}.docx"
    output_path = Path(settings.outputs_dir) / filename
    doc.save(str(output_path))
    return filename
